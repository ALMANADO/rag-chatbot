import streamlit as st
import os
import tempfile
import sys
import zipfile
import xml.etree.ElementTree as ET

# SQLite fix for Streamlit Cloud (Chroma sometimes needs this)
try:
    __import__("pysqlite3")
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass

from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument

from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_groq import ChatGroq
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import FakeEmbeddings
from langchain_core.prompts import ChatPromptTemplate


# -----------------------------
# DOCX Helpers (rich extraction)
# -----------------------------
def _extract_footnotes_from_docx(docx_path: str) -> str:
    try:
        with zipfile.ZipFile(docx_path) as z:
            if "word/footnotes.xml" not in z.namelist():
                return ""
            xml_bytes = z.read("word/footnotes.xml")
            root = ET.fromstring(xml_bytes)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            texts = [node.text for node in root.findall(".//w:t", ns) if node.text]
            return " ".join(texts).strip()
    except Exception:
        return ""


def load_docx_rich(docx_path: str, source_name: str = None) -> list[LCDocument]:
    doc = DocxDocument(docx_path)
    parts = []

    # Headers & Footers
    for si, section in enumerate(doc.sections):
        header_texts = [p.text.strip() for p in section.header.paragraphs if p.text and p.text.strip()]
        if header_texts:
            parts.append(f"\n[HEADER section {si + 1}]\n" + "\n".join(header_texts))

        footer_texts = [p.text.strip() for p in section.footer.paragraphs if p.text and p.text.strip()]
        if footer_texts:
            parts.append(f"\n[FOOTER section {si + 1}]\n" + "\n".join(footer_texts))

    # Paragraphs (mark headings)
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = (p.style.name if p.style else "").lower()
        if "heading" in style_name:
            parts.append(f"\n[HEADING] {text}\n")
        else:
            parts.append(text)

    # Tables
    for ti, table in enumerate(doc.tables):
        table_lines = [f"\n[TABLE {ti + 1}]"]
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = " ".join(t.strip() for t in cell.text.splitlines() if t.strip()).strip()
                row_cells.append(cell_text)
            table_lines.append(" | ".join(row_cells))
        parts.append("\n".join(table_lines))

    # Footnotes (best-effort)
    footnotes_text = _extract_footnotes_from_docx(docx_path)
    if footnotes_text:
        parts.append(f"\n[FOOTNOTES]\n{footnotes_text}")

    full_text = "\n".join(parts).strip()

    return [
        LCDocument(
            page_content=full_text,
            metadata={
                "source": source_name or os.path.basename(docx_path),
                "file_type": "docx",
            },
        )
    ]


# -----------------------------
# Streamlit UI Setup
# -----------------------------
st.set_page_config(page_title="RAG Chatbot", page_icon="🤖", layout="wide")

st.markdown(
    """
<style>
:root{
  --bg0:#05060a;
  --bg1:#0b0d14;
  --text:#e9eef7;
  --muted:rgba(233,238,247,.68);
  --muted2:rgba(233,238,247,.45);
  --glass:rgba(255,255,255,.06);
  --stroke:rgba(255,255,255,.12);
  --accent:#8b5cf6;
  --accent2:#22d3ee;
  --r-xl:26px;
  --r-lg:18px;
}

html, body, [data-testid="stAppViewContainer"]{
  color-scheme: dark !important;
}

html, body{
  height:100% !important;
  overflow-y:auto !important;
  background: var(--bg0) !important;
  color: var(--text) !important;
}

[data-testid="stAppViewContainer"]{
  height:100vh !important;
  overflow-y:auto !important;
  overflow-x:hidden !important;
  background:
    radial-gradient(1200px 600px at 50% 20%, rgba(139,92,246,.16), transparent 55%),
    radial-gradient(900px 520px at 85% 75%, rgba(34,211,238,.12), transparent 55%),
    linear-gradient(180deg, var(--bg0), var(--bg1)) !important;
  color: var(--text) !important;
  position:relative;
}

/* overlays */
[data-testid="stAppViewContainer"]::before{
  content:"";
  position:fixed;
  inset:0;
  pointer-events:none;
  z-index:0;
  background:
    radial-gradient(1px 1px at 12% 18%, rgba(255,255,255,.40) 50%, transparent 52%),
    radial-gradient(1px 1px at 22% 65%, rgba(255,255,255,.26) 50%, transparent 52%),
    radial-gradient(1px 1px at 35% 30%, rgba(255,255,255,.20) 50%, transparent 52%),
    radial-gradient(1px 1px at 44% 78%, rgba(255,255,255,.16) 50%, transparent 52%),
    radial-gradient(1px 1px at 58% 22%, rgba(255,255,255,.24) 50%, transparent 52%),
    radial-gradient(1px 1px at 70% 55%, rgba(255,255,255,.18) 50%, transparent 52%),
    radial-gradient(1px 1px at 82% 35%, rgba(255,255,255,.22) 50%, transparent 52%),
    radial-gradient(1px 1px at 90% 82%, rgba(255,255,255,.14) 50%, transparent 52%),
    linear-gradient(rgba(255,255,255,.04) 1px, transparent 1px),
    linear-gradient(90deg, rgba(255,255,255,.04) 1px, transparent 1px);
  background-size:auto,auto,auto,auto,auto,auto,auto,auto,44px 44px,44px 44px;
  opacity:.55;
}

[data-testid="stAppViewContainer"]::after{
  content:"";
  position:fixed;
  inset:0;
  pointer-events:none;
  z-index:0;
  background:radial-gradient(900px 520px at 50% 30%, transparent 40%, rgba(0,0,0,.55) 85%);
}

[data-testid="stAppViewContainer"] > *{
  position:relative;
  z-index:1;
}

header[data-testid="stHeader"]{
  background:transparent !important;
}

/* give bottom space so sticky input doesn't overlay content */
section.main > div{
  max-width:1050px;
  padding-top:1.2rem;
  padding-bottom:7rem;
}

section[data-testid="stSidebar"]{
  background:rgba(10,12,18,.65) !important;
  border-right:1px solid rgba(255,255,255,.06) !important;
  backdrop-filter:blur(14px);
  overflow-y:auto !important;
}

section[data-testid="stSidebar"] *{
  color:var(--text) !important;
}

h1{
  text-align:center;
  font-weight:650;
  letter-spacing:-.02em;
  margin-bottom:.2rem;
  color:var(--text) !important;
  text-shadow:0 0 24px rgba(255,255,255,.08), 0 0 45px rgba(139,92,246,.12);
}

h2{
  text-align:center;
  font-weight:500;
  color:var(--muted) !important;
}

.stCaption, [data-testid="stCaptionContainer"]{
  text-align:center !important;
  color:var(--muted) !important;
}

/* buttons */
.stButton > button{
  border-radius:999px !important;
  border:1px solid rgba(255,255,255,.10) !important;
  background:rgba(255,255,255,.06) !important;
  color:var(--text) !important;
  box-shadow:0 10px 26px rgba(0,0,0,.35) !important;
  transition:transform .15s ease, box-shadow .15s ease, border-color .15s ease;
}

.stButton > button:hover{
  transform:translateY(-1px);
  border-color:rgba(139,92,246,.35) !important;
  box-shadow:0 14px 32px rgba(0,0,0,.45), 0 0 20px rgba(139,92,246,.12) !important;
}

.stButton > button[kind="primary"]{
  background:linear-gradient(90deg, rgba(139,92,246,.22), rgba(34,211,238,.16)) !important;
  border:1px solid rgba(139,92,246,.35) !important;
}

/* uploader */
[data-testid="stFileUploaderDropzone"]{
  background:rgba(255,255,255,.05) !important;
  border:1px dashed rgba(255,255,255,.18) !important;
  border-radius:var(--r-xl) !important;
  padding:14px 14px !important;
}

[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzone"] p,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] label{
  color:rgba(233,238,247,.70) !important;
}

[data-testid="stFileUploaderDropzone"] button{
  border-radius:999px !important;
  border:1px solid rgba(255,255,255,.14) !important;
  background:rgba(255,255,255,.08) !important;
  color:rgba(233,238,247,.92) !important;
  padding:8px 14px !important;
  font-weight:600 !important;
  box-shadow:0 10px 22px rgba(0,0,0,.35) !important;
  margin-top:10px !important;
}

/* chat bubbles */
div[data-testid="stChatMessage"]{
  background:linear-gradient(180deg, rgba(255,255,255,.07), rgba(255,255,255,.04)) !important;
  border:1px solid rgba(255,255,255,.08) !important;
  border-radius:var(--r-xl) !important;
  box-shadow:0 10px 30px rgba(0,0,0,.35), 0 0 0 1px rgba(255,255,255,.03) inset !important;
  backdrop-filter:blur(12px);
}

div[role="img"][aria-label="assistant"] + div > div{
  background:linear-gradient(180deg, rgba(139,92,246,.16), rgba(255,255,255,.05)) !important;
  border:1px solid rgba(139,92,246,.22) !important;
}

/* sticky bottom input bar */
[data-testid="stBottom"]{
  position:sticky !important;
  bottom:0 !important;
  z-index:9999 !important;
  background:linear-gradient(180deg, rgba(5,6,10,0), rgba(5,6,10,.80) 35%, rgba(5,6,10,.92)) !important;
  backdrop-filter:blur(12px);
  padding-top:10px;
}

/* chat input base */
[data-testid="stChatInput"],
[data-testid="stChatInput"] form,
[data-testid="stChatInput"] > div{
  background:transparent !important;
}

/* textarea + input pill */
[data-testid="stChatInput"] textarea{
  background:rgba(255,255,255,.06) !important;
  border:1px solid rgba(255,255,255,.12) !important;
  color:rgba(233,238,247,.95) !important;
  border-radius:999px !important;
  min-height:52px !important;
  box-shadow:0 10px 24px rgba(0,0,0,.35) !important;
  padding:14px 78px 14px 18px !important;
}

[data-testid="stChatInput"] textarea::placeholder{
  color:rgba(233,238,247,.45) !important;
}

/* send button: aligned + centered */
[data-testid="stChatInput"] form{
  position:relative !important;
}

button[data-testid="stChatInputSubmitButton"]{
  position:absolute !important;
  right:14px !important;
  top:50% !important;
  transform:translateY(-50%) !important;
  margin:0 !important;
  height:44px !important;
  width:44px !important;
  min-width:44px !important;
  border-radius:999px !important;
  border:1px solid rgba(255,255,255,.10) !important;
  background:rgba(255,255,255,.06) !important;
  display:flex !important;
  align-items:center !important;
  justify-content:center !important;
  z-index:10 !important;
  box-shadow:0 10px 22px rgba(0,0,0,.30) !important;
}

button[data-testid="stChatInputSubmitButton"] svg{
  fill:rgba(233,238,247,.88) !important;
}

button[data-testid="stChatInputSubmitButton"]:hover{
  border-color:rgba(34,211,238,.35) !important;
  box-shadow:0 14px 28px rgba(0,0,0,.40), 0 0 18px rgba(34,211,238,.10) !important;
}

/* tooltip/popover fix (question-mark help etc.) */
div[data-baseweb="tooltip"],
div[data-baseweb="popover"]{
  z-index:100000 !important;
}

div[data-baseweb="tooltip"] > div,
div[data-baseweb="popover"] > div{
  background:rgba(12,14,22,.96) !important;
  color:rgba(233,238,247,.92) !important;
  border:1px solid rgba(255,255,255,.10) !important;
  border-radius:12px !important;
  box-shadow:0 18px 45px rgba(0,0,0,.55) !important;
  backdrop-filter:blur(12px);
  max-width:320px !important;
}

div[data-baseweb="tooltip"] * ,
div[data-baseweb="popover"] *{
  color:rgba(233,238,247,.92) !important;
  background:transparent !important;
}

/* scrollbar styling (page + sidebar + message container) */
[data-testid="stAppViewContainer"]::-webkit-scrollbar,
section[data-testid="stSidebar"]::-webkit-scrollbar,
div[data-testid="stVerticalBlock"]::-webkit-scrollbar{
  width:10px;
}

[data-testid="stAppViewContainer"]::-webkit-scrollbar-track,
section[data-testid="stSidebar"]::-webkit-scrollbar-track,
div[data-testid="stVerticalBlock"]::-webkit-scrollbar-track{
  background:rgba(255,255,255,.04);
  border-radius:999px;
}

[data-testid="stAppViewContainer"]::-webkit-scrollbar-thumb,
section[data-testid="stSidebar"]::-webkit-scrollbar-thumb,
div[data-testid="stVerticalBlock"]::-webkit-scrollbar-thumb{
  background:linear-gradient(180deg, rgba(139,92,246,.55), rgba(34,211,238,.35));
  border:2px solid rgba(0,0,0,.35);
  border-radius:999px;
}

[data-testid="stAppViewContainer"],
section[data-testid="stSidebar"],
div[data-testid="stVerticalBlock"]{
  scrollbar-width:thin;
  scrollbar-color:rgba(139,92,246,.60) rgba(255,255,255,.05);
}

a{ color:rgba(34,211,238,.95) !important; }
</style>
""",
    unsafe_allow_html=True,
)

st.title("🤖 RAG Chatbot with Groq")
st.header("Upload your documents and chat with them using ultra-fast LPU inference.")


# Session state init
if "messages" not in st.session_state:
    st.session_state.messages = []
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None


# Sidebar
with st.sidebar:
    st.header("📂 Document Management")
    uploaded_files = st.file_uploader(
        "Upload PDF, TXT, or DOCX files",
        accept_multiple_files=True,
        type=["pdf", "txt", "docx"],
        help="Drag and drop files here. Limit 200MB per file.",
    )

    col1, col2 = st.columns(2)
    with col1:
        process_btn = st.button("🚀 Process", type="primary")
    with col2:
        clear_btn = st.button("🗑️ Clear Chat")

    if clear_btn:
        st.session_state.messages = []
        st.rerun()

    if process_btn:
        if uploaded_files:
            with st.status("Processing documents...", expanded=True) as status:
                st.write("Loading files...")
                docs = []

                for uploaded_file in uploaded_files:
                    suffix = os.path.splitext(uploaded_file.name)[1].lower()
                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                        tmp.write(uploaded_file.getvalue())
                        tmp_path = tmp.name

                    try:
                        if suffix == ".pdf":
                            loader = PyPDFLoader(tmp_path)
                            docs.extend(loader.load())
                        elif suffix == ".docx":
                            docs.extend(load_docx_rich(tmp_path, source_name=uploaded_file.name))
                        else:
                            loader = TextLoader(tmp_path, encoding="utf-8")
                            docs.extend(loader.load())
                    finally:
                        if os.path.exists(tmp_path):
                            os.unlink(tmp_path)

                st.write("Splitting text into chunks...")
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                splits = text_splitter.split_documents(docs)

                st.write("Generating embeddings...")
                embeddings = FakeEmbeddings(size=384)

                st.write("Building vector index...")
                st.session_state.vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)
                st.session_state.retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": 4})

                status.update(label="✅ Processing Complete!", state="complete", expanded=False)
                st.success(f"Indexed {len(docs)} documents ({len(splits)} chunks).")
        else:
            st.warning("Please upload files first.")

    st.divider()
    st.info("Powered by Groq LPU Inference | Built with LangChain & Streamlit")


# -----------------------------
# Scrollable Chat History Area
# -----------------------------
chat_box = st.container(height=560, border=False)
with chat_box:
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])


# -----------------------------
# Chat Input
# -----------------------------
if prompt := st.chat_input("Ask a question about your documents..."):
    st.session_state.messages.append({"role": "user", "content": prompt})

    if st.session_state.retriever is None:
        response = "⚠️ Please upload and process documents in the sidebar first."
    else:
        with st.spinner("Analyzing context..."):
            try:
                groq_api_key = os.getenv("GROQ_API_KEY")
                if not groq_api_key:
                    response = "❌ **Error**: `GROQ_API_KEY` is missing. Please add it to Streamlit Secrets."
                else:
                    llm = ChatGroq(
                        model="llama-3.1-8b-instant",
                        temperature=0.0,
                        api_key=groq_api_key,
                    )

                    system_prompt = (
                        "You are a helpful AI assistant. "
                        "Use the provided context to answer the user's question accurately. "
                        "If the answer isn't in the context, say you don't know based on the documents. "
                        "Keep your response professional and helpful.\n\n"
                        "Context:\n{context}"
                    )

                    prompt_template = ChatPromptTemplate.from_messages(
                        [("system", system_prompt), ("human", "{input}")]
                    )

                    question_answer_chain = create_stuff_documents_chain(llm, prompt_template)
                    rag_chain = create_retrieval_chain(st.session_state.retriever, question_answer_chain)

                    result = rag_chain.invoke({"input": prompt})
                    response = result["answer"]

            except Exception as e:
                response = f"❌ **Error**: {str(e)}"

    st.session_state.messages.append({"role": "assistant", "content": response})
    st.rerun()
